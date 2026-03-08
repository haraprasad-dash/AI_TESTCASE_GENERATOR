"""
Document parsing service for extracting text from PDF, DOCX, and images.
"""
import os
from pathlib import Path
from typing import Optional, Dict, Any
from dataclasses import dataclass

import pytesseract
from PIL import Image

# Configure paths for tools installed via Scoop (no admin install)
_SCOOP_SHIMS = os.path.join(os.path.expanduser("~"), "scoop", "shims")
_SCOOP_TESSERACT = os.path.join(_SCOOP_SHIMS, "tesseract.exe")
if os.path.exists(_SCOOP_TESSERACT):
    pytesseract.pytesseract.tesseract_cmd = _SCOOP_TESSERACT

_SCOOP_POPPLER_BIN = os.path.join(
    os.path.expanduser("~"), "scoop", "apps", "poppler", "current", "bin"
)
_POPPLER_PATH = _SCOOP_POPPLER_BIN if os.path.exists(_SCOOP_POPPLER_BIN) else None


@dataclass
class ExtractedContent:
    """Result of document text extraction."""
    filename: str
    content_type: str
    text: str
    page_count: Optional[int] = None
    error: Optional[str] = None


class DocumentParserError(Exception):
    """Base exception for document parsing errors."""
    pass


class UnsupportedFileError(DocumentParserError):
    """File type not supported."""
    pass


class ExtractionError(DocumentParserError):
    """Failed to extract text from file."""
    pass


class DocumentParser:
    """Parser for extracting text from various document formats."""
    
    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "image/png": "image",
        "image/jpeg": "image",
        "image/jpg": "image"
    }
    
    MAX_FILE_SIZE_MB = 20
    
    def __init__(self):
        self.ocr_available = self._check_ocr()
    
    def _check_ocr(self) -> bool:
        """Check if Tesseract OCR is available."""
        try:
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            return False
    
    async def parse_file(self, file_path: str, content_type: str) -> ExtractedContent:
        """
        Parse a file and extract text content.
        
        Args:
            file_path: Path to the file
            content_type: MIME type of the file
            
        Returns:
            ExtractedContent with text and metadata
            
        Raises:
            UnsupportedFileError: If file type not supported
            ExtractionError: If text extraction fails
        """
        filename = os.path.basename(file_path)
        
        # Check file size
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > self.MAX_FILE_SIZE_MB:
            raise ExtractionError(f"File too large: {file_size_mb:.1f}MB (max {self.MAX_FILE_SIZE_MB}MB)")
        
        # Determine file type
        doc_type = self.SUPPORTED_TYPES.get(content_type)
        if not doc_type:
            raise UnsupportedFileError(f"Unsupported file type: {content_type}")
        
        try:
            if doc_type == "pdf":
                return await self._parse_pdf(file_path, content_type)
            elif doc_type == "docx":
                return await self._parse_docx(file_path, content_type)
            elif doc_type == "image":
                return await self._parse_image(file_path, content_type)
            else:
                raise UnsupportedFileError(f"Unknown document type: {doc_type}")
                
        except Exception as e:
            raise ExtractionError(f"Failed to extract text: {str(e)}")
    
    async def _parse_pdf(self, file_path: str, content_type: str) -> ExtractedContent:
        """Extract text from PDF file."""
        try:
            import PyPDF2
        except ImportError:
            raise ExtractionError("PyPDF2 not installed")
        
        text_parts = []
        page_count = 0
        
        with open(file_path, 'rb') as file:
            try:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"\n--- Page {page_num + 1} ---\n")
                        text_parts.append(page_text)
                
                extracted_text = "\n".join(text_parts)
                
                # If no text extracted, try OCR
                if not extracted_text.strip() and self.ocr_available:
                    extracted_text = await self._ocr_pdf(file_path)
                
                return ExtractedContent(
                    filename=os.path.basename(file_path),
                    content_type=content_type,
                    text=extracted_text,
                    page_count=page_count
                )
                
            except Exception as e:
                # Try OCR as fallback
                if self.ocr_available:
                    text = await self._ocr_pdf(file_path)
                    return ExtractedContent(
                        filename=os.path.basename(file_path),
                        content_type=content_type,
                        text=text,
                        page_count=page_count
                    )
                raise ExtractionError(f"PDF parsing failed: {e}")
    
    async def _parse_docx(self, file_path: str, content_type: str) -> ExtractedContent:
        """Extract text from DOCX file."""
        try:
            from docx import Document
        except ImportError:
            raise ExtractionError("python-docx not installed")
        
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            extracted_text = "\n\n".join(paragraphs)
            
            return ExtractedContent(
                filename=os.path.basename(file_path),
                content_type=content_type,
                text=extracted_text,
                page_count=None
            )
            
        except Exception as e:
            raise ExtractionError(f"DOCX parsing failed: {e}")
    
    async def _parse_image(self, file_path: str, content_type: str) -> ExtractedContent:
        """Extract text from image using OCR."""
        if not self.ocr_available:
            raise ExtractionError("OCR not available. Install tesseract-ocr.")
        
        try:
            image = Image.open(file_path)
            
            # Preprocess for better OCR
            image = self._preprocess_image(image)
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            return ExtractedContent(
                filename=os.path.basename(file_path),
                content_type=content_type,
                text=text.strip(),
                page_count=1
            )
            
        except Exception as e:
            raise ExtractionError(f"Image OCR failed: {e}")
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """Preprocess image for better OCR results."""
        # Convert to grayscale
        if image.mode != 'L':
            image = image.convert('L')
        
        # Increase contrast
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)
        
        return image
    
    async def _ocr_pdf(self, file_path: str) -> str:
        """Perform OCR on PDF pages."""
        if not self.ocr_available:
            return ""
        
        try:
            from pdf2image import convert_from_path
        except ImportError:
            return ""
        
        text_parts = []
        
        try:
            images = convert_from_path(file_path, dpi=200, poppler_path=_POPPLER_PATH)
            
            for i, image in enumerate(images):
                text = pytesseract.image_to_string(image)
                if text.strip():
                    text_parts.append(f"\n--- Page {i + 1} ---\n")
                    text_parts.append(text)
            
            return "\n".join(text_parts)
            
        except Exception as e:
            raise ExtractionError(f"PDF OCR failed: {e}")


# Global parser instance
_document_parser: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """Get or create document parser instance."""
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParser()
    return _document_parser
